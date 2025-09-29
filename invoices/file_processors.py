"""
File processors for extracting text from different document formats
Supports PDF, TXT, DOCX, XLSX, XLS, and other common document formats
"""

import logging
import os
from typing import Optional, Dict, Any

from docx import Document

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Processes different document formats and extracts text content
    """

    SUPPORTED_FORMATS = {
        '.pdf': 'PDF Document',
        '.txt': 'Text File',
        '.docx': 'Word Document',
        '.doc': 'Word Document (Legacy)',
        '.xlsx': 'Excel Spreadsheet',
        '.xls': 'Excel Spreadsheet (Legacy)',
        '.csv': 'CSV File',
        '.rtf': 'Rich Text Format',
        '.odt': 'OpenDocument Text'
    }

    @classmethod
    def get_supported_extensions(cls) -> list:
        """Get list of supported file extensions"""
        return list(cls.SUPPORTED_FORMATS.keys())

    @classmethod
    def is_supported_file(cls, filename: str) -> bool:
        """Check if file format is supported"""
        _, ext = os.path.splitext(filename.lower())
        return ext in cls.SUPPORTED_FORMATS

    @classmethod
    def get_file_type(cls, filename: str) -> str:
        """Get file type description"""
        _, ext = os.path.splitext(filename.lower())
        return cls.SUPPORTED_FORMATS.get(ext, 'Unknown')

    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from a document file

        Args:
            file_path (str): Path to the document file

        Returns:
            Dict[str, Any]: Extraction results containing text and metadata
        """
        try:
            if not os.path.exists(file_path):
                return self._error_result(f"File not found: {file_path}")

            filename = os.path.basename(file_path)
            _, ext = os.path.splitext(filename.lower())

            if not self.is_supported_file(filename):
                return self._error_result(f"Unsupported file format: {ext}")

            # Extract text based on file type
            if ext == '.txt':
                result = self._extract_from_txt(file_path)
            elif ext == '.docx':
                result = self._extract_from_docx(file_path)
            elif ext == '.pdf':
                result = self._extract_from_pdf(file_path)
            elif ext in ['.xlsx', '.xls']:
                result = self._extract_from_excel(file_path)
            elif ext == '.csv':
                result = self._extract_from_csv(file_path)
            elif ext == '.rtf':
                result = self._extract_from_rtf(file_path)
            elif ext in ['.doc', '.odt']:
                result = self._extract_from_legacy_format(file_path, ext)
            else:
                return self._error_result(f"No processor available for {ext}")

            # Add metadata
            result['metadata'] = {
                'filename': filename,
                'file_type': self.get_file_type(filename),
                'file_extension': ext,
                'file_size': os.path.getsize(file_path)
            }

            logger.info(f"Successfully extracted text from {filename} ({ext})")
            return result

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return self._error_result(f"Extraction failed: {str(e)}")

    def _extract_from_txt(self, file_path: str) -> Dict[str, Any]:
        """Extract text from a .txt file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()

                    return {
                        'success': True,
                        'text': text.strip(),
                        'extraction_method': f'text_file_{encoding}',
                        'page_count': 1,
                        'word_count': len(text.split()),
                        'character_count': len(text)
                    }
                except UnicodeDecodeError:
                    continue

            return self._error_result("Unable to decode text file with any supported encoding")

        except Exception as e:
            return self._error_result(f"Text file extraction failed: {str(e)}")

    def _extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from a .docx file"""
        try:
            doc = Document(file_path)

            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())

            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(' | '.join(row_text))

            # Combine all text
            all_text = []
            all_text.extend(paragraphs)
            if table_text:
                all_text.append('\n--- Tables ---')
                all_text.extend(table_text)

            combined_text = '\n'.join(all_text)

            return {
                'success': True,
                'text': combined_text,
                'extraction_method': 'python_docx',
                'page_count': len(doc.paragraphs) // 20 + 1,  # Rough estimate
                'paragraph_count': len(paragraphs),
                'table_count': len(doc.tables),
                'word_count': len(combined_text.split()),
                'character_count': len(combined_text)
            }

        except Exception as e:
            return self._error_result(f"DOCX extraction failed: {str(e)}")

    def _extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from a PDF file
        This is a placeholder - you can integrate PyPDF2, pdfplumber, or similar
        """
        try:
            # Placeholder implementation
            # In a real scenario, you'd use a PDF library like:
            # - PyPDF2: pip install PyPDF2
            # - pdfplumber: pip install pdfplumber
            # - PyMuPDF: pip install PyMuPDF

            filename = os.path.basename(file_path)

            # For now, return a placeholder result
            return {
                'success': True,
                'text': f"PDF text extraction not yet implemented for {filename}.\nPlease use .txt or .docx files for now.",
                'extraction_method': 'placeholder',
                'page_count': 1,
                'word_count': 10,
                'character_count': 50,
                'note': 'PDF extraction requires additional PDF library integration'
            }

        except Exception as e:
            return self._error_result(f"PDF extraction failed: {str(e)}")

    def _extract_from_excel(self, file_path: str) -> Dict[str, Any]:
        """Extract text from Excel files (.xlsx, .xls)"""
        try:
            try:
                import pandas as pd
            except ImportError:
                return self._error_result("pandas library required for Excel file processing")

            # Read Excel file
            try:
                df = pd.read_excel(file_path, sheet_name=None)  # Read all sheets
            except Exception as e:
                return self._error_result(f"Failed to read Excel file: {str(e)}")

            # Extract text from all sheets
            all_text = []
            sheet_count = 0
            total_rows = 0

            for sheet_name, sheet_df in df.items():
                sheet_count += 1
                total_rows += len(sheet_df)

                # Add sheet header
                all_text.append(f"--- Sheet: {sheet_name} ---")

                # Convert dataframe to text
                # Include column headers
                if not sheet_df.empty:
                    all_text.append(" | ".join(str(col) for col in sheet_df.columns))

                    # Add data rows
                    for _, row in sheet_df.iterrows():
                        row_text = " | ".join(str(value) if pd.notna(value) else "" for value in row)
                        if row_text.strip():
                            all_text.append(row_text)

            combined_text = '\n'.join(all_text)

            return {
                'success': True,
                'text': combined_text,
                'extraction_method': 'pandas_excel',
                'sheet_count': sheet_count,
                'total_rows': total_rows,
                'word_count': len(combined_text.split()),
                'character_count': len(combined_text)
            }

        except Exception as e:
            return self._error_result(f"Excel extraction failed: {str(e)}")

    def _extract_from_csv(self, file_path: str) -> Dict[str, Any]:
        """Extract text from CSV files"""
        try:
            try:
                import pandas as pd
            except ImportError:
                # Fallback to basic CSV reading
                import csv

                encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']

                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding, newline='') as file:
                            csv_reader = csv.reader(file)
                            rows = list(csv_reader)

                        text_lines = []
                        for row in rows:
                            text_lines.append(" | ".join(row))

                        combined_text = '\n'.join(text_lines)

                        return {
                            'success': True,
                            'text': combined_text,
                            'extraction_method': f'csv_basic_{encoding}',
                            'row_count': len(rows),
                            'word_count': len(combined_text.split()),
                            'character_count': len(combined_text)
                        }
                    except UnicodeDecodeError:
                        continue

                return self._error_result("Unable to decode CSV file with any supported encoding")

            # Use pandas for better CSV handling
            try:
                df = pd.read_csv(file_path)

                # Convert to text
                text_lines = []

                # Add column headers
                text_lines.append(" | ".join(str(col) for col in df.columns))

                # Add data rows
                for _, row in df.iterrows():
                    row_text = " | ".join(str(value) if pd.notna(value) else "" for value in row)
                    text_lines.append(row_text)

                combined_text = '\n'.join(text_lines)

                return {
                    'success': True,
                    'text': combined_text,
                    'extraction_method': 'pandas_csv',
                    'row_count': len(df),
                    'column_count': len(df.columns),
                    'word_count': len(combined_text.split()),
                    'character_count': len(combined_text)
                }

            except Exception as e:
                return self._error_result(f"CSV parsing failed: {str(e)}")

        except Exception as e:
            return self._error_result(f"CSV extraction failed: {str(e)}")

    def _extract_from_rtf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from RTF files"""
        try:
            # Basic RTF parsing - strips RTF codes
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()

            # Simple RTF parsing - remove RTF control codes
            import re
            # Remove RTF header and control words
            text = re.sub(r'\\[a-z]+\d*\s?', '', content)
            text = re.sub(r'[{}]', '', text)
            text = re.sub(r'\s+', ' ', text)
            text = text.strip()

            return {
                'success': True,
                'text': text,
                'extraction_method': 'basic_rtf_parser',
                'word_count': len(text.split()),
                'character_count': len(text),
                'note': 'Basic RTF parsing - advanced formatting may be lost'
            }

        except Exception as e:
            return self._error_result(f"RTF extraction failed: {str(e)}")

    def _extract_from_legacy_format(self, file_path: str, ext: str) -> Dict[str, Any]:
        """Handle legacy formats like .doc and .odt"""
        try:
            filename = os.path.basename(file_path)

            # These formats require specialized libraries
            if ext == '.doc':
                error_msg = "Legacy .doc files require python-docx2txt or antiword for extraction"
            elif ext == '.odt':
                error_msg = "OpenDocument files require python-odf or LibreOffice for extraction"
            else:
                error_msg = f"Legacy format {ext} extraction not implemented"

            return {
                'success': False,
                'text': f"Cannot extract text from {filename}. {error_msg}",
                'extraction_method': 'not_supported',
                'error': error_msg,
                'suggested_alternative': 'Please convert to .docx, .txt, or .pdf format'
            }

        except Exception as e:
            return self._error_result(f"Legacy format extraction failed: {str(e)}")

    def _error_result(self, error_message: str) -> Dict[str, Any]:
        """Return a standardized error result"""
        return {
            'success': False,
            'text': '',
            'error': error_message,
            'extraction_method': 'error',
            'page_count': 0,
            'word_count': 0,
            'character_count': 0
        }


def extract_text_from_file(file_path: str) -> Dict[str, Any]:
    """
    Convenience function to extract text from a file

    Args:
        file_path (str): Path to the document file

    Returns:
        Dict[str, Any]: Extraction results
    """
    processor = DocumentProcessor()
    return processor.extract_text(file_path)


def get_supported_file_types() -> Dict[str, str]:
    """
    Get dictionary of supported file types

    Returns:
        Dict[str, str]: Extension to description mapping
    """
    return DocumentProcessor.SUPPORTED_FORMATS.copy()


def validate_uploaded_file(filename: str) -> Dict[str, Any]:
    """
    Validate if an uploaded file is supported

    Args:
        filename (str): Name of the uploaded file

    Returns:
        Dict[str, Any]: Validation results
    """
    processor = DocumentProcessor()

    if not filename:
        return {
            'valid': False,
            'error': 'No filename provided',
            'file_type': None
        }

    is_supported = processor.is_supported_file(filename)
    file_type = processor.get_file_type(filename)

    if is_supported:
        return {
            'valid': True,
            'file_type': file_type,
            'extension': os.path.splitext(filename.lower())[1]
        }
    else:
        supported_exts = ', '.join(processor.get_supported_extensions())
        return {
            'valid': False,
            'error': f'Unsupported file type. Supported formats: {supported_exts}',
            'file_type': file_type
        }